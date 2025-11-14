"""
agent_fastapi.py
-------------------
FastAPI + LangGraph Content Creation Agent with Streaming Support

Features:
- REST API endpoints
- Real-time streaming responses
- Markdown formatted output
- CORS enabled for frontend integration
- Async processing
- Tavily Web Search with Image Support

Run: uvicorn agent_fastapi:app --reload --host 0.0.0.0 --port 8000
"""

import dotenv
dotenv.load_dotenv()    # very important. 

import httpx

import os
import json
import re
import asyncio
import io
import base64
import resend
from docx import Document
from bs4 import BeautifulSoup
from typing import TypedDict, List, Dict, Any, Optional
from datetime import datetime
from concurrent.futures import ThreadPoolExecutor

from fastapi import FastAPI, HTTPException, BackgroundTasks, Header
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import StreamingResponse, JSONResponse
from pydantic import BaseModel, Field
from langchain_ollama import OllamaLLM
from langchain_tavily import TavilySearch # 1. IMPORT TAVILY
import secrets

# -------------------------------
# FastAPI App Setup
# -------------------------------
app = FastAPI(
    title="AI Content Creation Agent",
    description="LangGraph-powered content generation with Ollama and Tavily Search",
    version="1.1.0"
)

# Enable CORS for frontend
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],  # In production, specify your frontend URL
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# -------------------------------
# Authentication Setup
# -------------------------------
HARDCODED_USERS = {
    "admin": "creator123",
    "user": "password123"
}

active_tokens = set()

# -------------------------------
# LLM Setup
# -------------------------------
LLM_MODEL = os.getenv("OLLAMA_MODEL", "llama3")

llm = OllamaLLM(
    model=LLM_MODEL,
    temperature=0.1,
    num_predict=512,
    num_ctx=2048,
    top_p=0.9,
    top_k=40
)

llm_writer = OllamaLLM(
    model=LLM_MODEL,
    temperature=0.2,
    num_predict=1024,
    num_ctx=4096,  # Increased context for search results
    top_p=0.9,
    top_k=40
)

# -------------------------------
# Pydantic Models (Request/Response)
# -------------------------------
class ContentRequest(BaseModel):
    project_brief: str = Field(..., description="Brief description of the content project")
    audience: str = Field(default="general", description="Target audience")
    tone: str = Field(default="informative", description="Content tone")
    keywords: List[str] = Field(default=[], description="Keywords to include")
    content_type: str = Field(default="article", description="Type of content to generate")
    # 2. ADD WEB_SEARCH TO REQUEST MODEL
    web_search: int = Field(default=0, description="Number of web search results to include (0 to disable)")
    author_email: Optional[str] = Field(default=None, description="Author email for sending content")

class ContentResponse(BaseModel):
    status: str
    strategy: Optional[str] = None
    draft_content: Optional[str] = None
    quality_score: Optional[float] = None
    reviewer_comments: Optional[str] = None
    needs_escalation: bool = False
    final_content: Optional[str] = None
    timestamp: str

class StreamUpdate(BaseModel):
    stage: str
    content: str
    progress: int

class LoginRequest(BaseModel):
    username: str
    password: str

class LoginResponse(BaseModel):
    token: str
    message: str

# -------------------------------
# State Schema
# -------------------------------
class ContentState(TypedDict, total=False):
    project_brief: str
    audience: str
    tone: str
    keywords: List[str]
    content_type: str
    # 3. ADD WEB_SEARCH KEYS TO STATE
    web_search: int
    search_results: Optional[List[Dict[str, Any]]]
    strategy: str
    draft_content: str
    quality_score: float
    needs_escalation: bool
    reviewer_comments: str
    final_content: str


# -------------------------------
# Agent Nodes (Async versions)
# -------------------------------

# 4. ADD NEW WEB_SEARCH_ASYNC NODE
async def web_search_async(state: ContentState) -> ContentState:
    """Search the web for the project brief if requested."""
    
    max_results = state.get("web_search", 0)
    
    if max_results == 0:
        print("Skipping web search.")
        state["search_results"] = None
        return state
        
    print(f"Conducting web search for: {state.get('project_brief')}")
    
    loop = asyncio.get_event_loop()
    with ThreadPoolExecutor() as executor:
        # Instantiate TavilySearch, including images
        search_tool = TavilySearch(max_results=max_results, include_images=True)
        try:
            results = await loop.run_in_executor(
                executor,
                search_tool.invoke,
                state.get("project_brief")
            )
            state["search_results"] = results
            print(f"Found {len(results)} search results.")
        except Exception as e:
            print(f"Error during web search: {e}")
            state["search_results"] = None
            
    return state


async def content_strategist_async(state: ContentState) -> ContentState:
    """Generate content strategy"""
    brief = state.get("project_brief", "")
    audience = state.get("audience", "general")
    tone = state.get("tone", "informative")
    keywords = ", ".join(state.get("keywords", []))
    content_type = state.get("content_type", "article")
    
    # 5. UPDATE NODE TO USE SEARCH RESULTS
    search_results = state.get("search_results")
    search_context = "No web search results."
    if search_results:
        search_context = "## Web Search Context:\n"
        for i, result in enumerate(search_results.get('results', [])): # bug 
            search_context += f"Source {i+1} ({result.get('url')}):\n{result.get('content')}\n\n"

    prompt = f"""You are an advanced Content Strategist Agent.  
Your task is to analyze the brief, audience, tone, keywords, content_type, and optional web search context to produce a complete, structured strategy.

{search_context}

---

# INPUT DETAILS
**Content Type:** {content_type}  
**Topic / Brief:** {brief}  
**Audience:** {audience}  
**Tone:** {tone}  
**Keywords:** {keywords}  

---

# GLOBAL STRATEGY RULES
You MUST follow these rules regardless of content_type:

1. Produce a deeply structured, multi-layered strategy.
2. Use clear sections with markdown headings.
3. Make decisions — do NOT be vague.
4. Integrate insights from web search ONLY if relevant.
5. Tailor everything to the selected content_type.

---"""

    loop = asyncio.get_event_loop()
    with ThreadPoolExecutor() as executor:
        response = await loop.run_in_executor(executor, llm.invoke, prompt)
    
    state["strategy"] = response.strip()
    return state


async def content_writer_async(state: ContentState) -> ContentState:
    """Write the actual content"""
    strategy = state.get("strategy", "")
    tone = state.get("tone", "informative")
    keywords = ", ".join(state.get("keywords", []))
    content_type = state.get("content_type", "article")

    # 6. UPDATE NODE TO USE SEARCH RESULTS AND IMAGES
    search_results = state.get("search_results")
    search_context = "No web search results."
    if search_results:
        search_context = "## Web Search Context (for your reference):\n"
        for i, result in enumerate(search_results.get('results', [])):
            search_context += f"Source {i+1} ({result.get('url')}):\n{result.get('content')}\n"
            if result.get('images'):
                 search_context += f"Images for Source {i+1}: {', '.join(result.get('images'))}\n\n"

    prompt = f"""You are a Senior Content Writer Agent.  
Your task is to convert the strategist_output into fully finished content based on the selected content_type.

# STRATEGIST OUTPUT (REFERENCE)
{strategy}

---

# SEARCH CONTEXT (REFERENCE)
{search_context}

---

# INPUT DETAILS
**Content Type:** {content_type}  
**Tone:** {tone}  
**Keywords:** {keywords}  

---

# GLOBAL WRITING RULES
1. Follow strategist_output EXACTLY.  
2. Maintain the specified tone throughout.  
3. Expand ideas with clarity, depth, and logical progression.  
4. Use markdown headings (##, ###) where appropriate.  
5. Use examples, reasoning, and explanation where beneficial.  
6. Integrate keywords naturally.  
7. If images exist in search context, embed where relevant:  
   `![alt text](image_url)`  
8. Return ONLY the final content — do NOT add commentary.

---

# CONTENT-TYPE–SPECIFIC WRITING BEHAVIOR

## If content_type == "blogger"
Write:
### Full Markdown Article
- H1 Title  
- 4–7 H2 sections  
- Optional H3 subsections  
- Detailed paragraphs  
- Examples, data, insights  

### Include:
- Smooth transitions  
- Data-backed reasoning  
- Optional images (if found)  

---

## If content_type == "social_media"
Write:
### 2–3 Post Variations
Each variation must include:
- Strong hook  
- Body  
- CTA  

Platform rules must be followed:
- Twitter/X → concise  
- Instagram → spacing + emphasis  
- LinkedIn → structured insights  

---

## If content_type == "copywriter"
Write:
### Ad Copy
- Short (1–2 lines)  
- Medium (3–5 lines)  
- Long (paragraph style)

### Landing Page Section
- Hero line  
- Subheadline  
- CTA

### Optional:
- 6–10 taglines  

---

## If content_type == "newsletter"
Write:
### Complete Newsletter
- Warm intro  
- Main narrative  
- Insight section  
- Closing message  
- CTA or forward prompt  

Tone: conversational, editorial, clear.  

---

## If content_type == "podcaster"
Write:
### Full Podcast Script
- Opening lines  
- Segment-by-segment script  
- Natural spoken flow  
- Pacing cues  
- For interviews:  
  - Host questions  
  - Follow-ups  

---

## If content_type == "youtuber"
Write:
### Complete YouTube Script
- Hook  
- Scene-by-scene script  
- On-screen text suggestions  
- B-roll cues  
- Strong outro  

Keep sentences punchy and spoken-language friendly.  

---

# OUTPUT FORMAT (MANDATORY)
Return ONLY:

## Writer Output
<final polished content here>"""

    loop = asyncio.get_event_loop()
    with ThreadPoolExecutor() as executor:
        response = await loop.run_in_executor(executor, llm_writer.invoke, prompt)
    
    state["draft_content"] = response.strip()
    return state


async def content_reviewer_async(state: ContentState) -> ContentState:
    """Review content quality"""
    draft = state.get("draft_content", "")
    tone = state.get("tone", "informative")
    keywords = ", ".join(state.get("keywords", []))
    
    if not draft:
        state["quality_score"] = 0.0
        state["reviewer_comments"] = "No draft to review."
        state["needs_escalation"] = True
        return state

    prompt = f"""You are a Senior Content Quality Reviewer Agent.  
Your task is to evaluate the draft content rigorously and return ONLY a JSON object.

# DRAFT UNDER REVIEW (TRUNCATED)
{draft[:2000]}

---

# REVIEW CRITERIA
Evaluate the content against the following:

## 1. Structural Quality
- Does it follow strategist_output?
- Is the format correct for the content_type?
- Is the flow logical?

## 2. Clarity & Readability
- Straightforward sentences?
- Easy to follow?

## 3. Depth & Completeness
- Key points fully developed?
- Missing sections?

## 4. Tone Alignment
Compare against:  
Tone specified → {tone}  

## 5. Keyword Usage
Check natural inclusion of:  
{keywords}

## 6. Suitability for Content Type
Specific expectations:
- Blogger → full article with structure  
- Social media → concise, hook-based  
- Copywriter → persuasion-first  
- Newsletter → editorial clarity  
- Podcaster → spoken flow  
- Youtuber → script + cues  

---

# SCORE RULES
- 9–10 = Excellent  
- 7–8 = Good  
- 6 = Acceptable but needs light polishing  
- <6 = Needs Escalation (fail)

---

# OUTPUT FORMAT (STRICT)
Return ONLY this JSON object:

{{
  "score": <number 0–10>,
  "comments": "<specific, actionable feedback>"
}}

NO text outside the JSON.  
NO explanations.  
NO markdown."""

    loop = asyncio.get_event_loop()
    with ThreadPoolExecutor() as executor:
        review_text = await loop.run_in_executor(executor, llm.invoke, prompt)

    try:
        match = re.search(r"\{[^{}]*\}", review_text)
        if match:
            data = json.loads(match.group(0))
        else:
            data = {"score": 7.0, "comments": "Review completed"}
    except Exception:
        data = {"score": 7.0, "comments": review_text[:200]}

    state["quality_score"] = float(data.get("score", 7.0))
    state["reviewer_comments"] = data.get("comments", "No comments")
    state["needs_escalation"] = state["quality_score"] < 6.0
    return state


async def finalizer_async(state: ContentState) -> ContentState:
    """Finalize content or escalate"""
    if state.get("needs_escalation"):
        state["final_content"] = ""
        return state

    state["final_content"] = state.get("draft_content", "")
    return state


# -------------------------------
# Email Functionality
# -------------------------------
def create_docx_from_html(html_content: str, title: str = "Generated Content") -> io.BytesIO:
    """Create a DOCX file from HTML content (same as frontend download)"""
    from bs4 import BeautifulSoup
    
    doc = Document()
    doc.add_heading(title, 0)
    
    # Parse HTML content
    soup = BeautifulSoup(html_content, 'html.parser')
    
    # Process each element
    for element in soup.find_all(['h1', 'h2', 'h3', 'h4', 'h5', 'h6', 'p', 'ul', 'ol', 'li']):
        if element.name in ['h1', 'h2', 'h3', 'h4', 'h5', 'h6']:
            level = int(element.name[1])
            doc.add_heading(element.get_text().strip(), level=level)
        elif element.name == 'p':
            text = element.get_text().strip()
            if text:
                p = doc.add_paragraph()
                # Handle bold text
                for child in element.children:
                    if hasattr(child, 'name'):
                        if child.name == 'strong' or child.name == 'b':
                            p.add_run(child.get_text()).bold = True
                        else:
                            p.add_run(child.get_text())
                    else:
                        p.add_run(str(child))
        elif element.name == 'ul':
            for li in element.find_all('li'):
                doc.add_paragraph(li.get_text().strip(), style='List Bullet')
        elif element.name == 'ol':
            for li in element.find_all('li'):
                doc.add_paragraph(li.get_text().strip(), style='List Number')
    
    # Save to BytesIO
    docx_buffer = io.BytesIO()
    doc.save(docx_buffer)
    docx_buffer.seek(0)
    return docx_buffer


def create_docx_from_content(content: str, title: str = "Generated Content") -> io.BytesIO:
    """Create a DOCX file from markdown content with better formatting"""
    doc = Document()
    doc.add_heading(title, 0)
    
    # Enhanced markdown to docx conversion
    lines = content.split('\n')
    in_list = False
    
    for line in lines:
        line = line.strip()
        
        if not line:  # Empty line
            if in_list:
                in_list = False
            continue
            
        # Headers
        if line.startswith('# '):
            if in_list: in_list = False
            doc.add_heading(line[2:], level=1)
        elif line.startswith('## '):
            if in_list: in_list = False
            doc.add_heading(line[3:], level=2)
        elif line.startswith('### '):
            if in_list: in_list = False
            doc.add_heading(line[4:], level=3)
        elif line.startswith('#### '):
            if in_list: in_list = False
            doc.add_heading(line[5:], level=4)
        
        # Lists
        elif line.startswith('- ') or line.startswith('* '):
            doc.add_paragraph(line[2:], style='List Bullet')
            in_list = True
        elif line.startswith(('1. ', '2. ', '3. ', '4. ', '5. ', '6. ', '7. ', '8. ', '9. ')):
            doc.add_paragraph(line[3:], style='List Number')
            in_list = True
        
        # Bold text (simple **text** conversion)
        elif '**' in line:
            if in_list: in_list = False
            p = doc.add_paragraph()
            parts = line.split('**')
            for i, part in enumerate(parts):
                if i % 2 == 0:
                    p.add_run(part)
                else:
                    p.add_run(part).bold = True
        
        # Regular paragraph
        elif line:
            if in_list: in_list = False
            doc.add_paragraph(line)
    
    # Save to BytesIO
    docx_buffer = io.BytesIO()
    doc.save(docx_buffer)
    docx_buffer.seek(0)
    return docx_buffer


async def send_content_email_with_docx(recipient_email: str, docx_base64: str, project_brief: str):
    """Send generated content via email with pre-generated DOCX attachment"""
    # Get Resend API key from environment
    resend_api_key = os.getenv("RESEND_API_KEY")
    
    if not resend_api_key:
        raise Exception("RESEND_API_KEY not set. Get free API key from https://resend.com")
    
    # Set API key
    resend.api_key = resend_api_key
    
    filename = f"{project_brief[:30].replace(' ', '_')}_content.docx"
    
    # Send email using Resend with the exact DOCX from frontend
    loop = asyncio.get_event_loop()
    with ThreadPoolExecutor() as executor:
        await loop.run_in_executor(
            executor,
            _send_resend_email,
            recipient_email, project_brief, docx_base64, filename
        )


def _send_resend_email(recipient_email, project_brief, docx_base64, filename):
    """Send email using Resend API"""
    resend.Emails.send({
        "from": "Creator.ai <noreply@resend.dev>",  # Free domain provided by Resend
        "to": recipient_email,
        "subject": f"🚀 Your AI-Generated Content is Ready! - {project_brief[:40]}...",
        "html": f"""
        <div style="font-family: Arial, sans-serif; max-width: 600px; margin: 0 auto; padding: 20px; background-color: #f8f9fa;">
            <div style="background: linear-gradient(135deg, #667eea 0%, #764ba2 100%); padding: 30px; border-radius: 10px; text-align: center; margin-bottom: 20px;">
                <h1 style="color: white; margin: 0; font-size: 28px;">🤖 Creator.ai</h1>
                <p style="color: #e8f4fd; margin: 10px 0 0 0; font-size: 16px;">Your AI Content Creation Team</p>
            </div>
            
            <div style="background: white; padding: 30px; border-radius: 10px; box-shadow: 0 2px 10px rgba(0,0,0,0.1);">
                <h2 style="color: #333; margin-top: 0;">🎉 Your Content is Ready!</h2>
                
                <p style="color: #555; font-size: 16px; line-height: 1.6;">Hi there! 👋</p>
                
                <p style="color: #555; font-size: 16px; line-height: 1.6;">
                    Your AI-generated content has been successfully created by our multi-agent team! 
                    📝 Our Strategist, Writer, and Reviewer have worked together to craft high-quality content just for you.
                </p>
                
                <div style="background: #f8f9fa; padding: 20px; border-radius: 8px; margin: 20px 0; border-left: 4px solid #667eea;">
                    <p style="margin: 0; color: #333;"><strong>📋 Project Brief:</strong></p>
                    <p style="margin: 5px 0 0 0; color: #666; font-style: italic;">{project_brief}</p>
                </div>
                
                <p style="color: #555; font-size: 16px; line-height: 1.6;">
                    📎 Please find your professionally formatted content attached as a DOCX file, 
                    ready for immediate use!
                </p>
                
                <div style="text-align: center; margin: 30px 0;">
                    <div style="background: #e8f5e8; padding: 15px; border-radius: 8px; display: inline-block;">
                        <p style="margin: 0; color: #2d5a2d; font-weight: bold;">✨ What's Included:</p>
                        <p style="margin: 5px 0 0 0; color: #2d5a2d;">📊 Strategic Planning • ✍️ Professional Writing • 🔍 Quality Review</p>
                    </div>
                </div>
                
                <hr style="border: none; height: 1px; background: #eee; margin: 30px 0;">
                
                <p style="color: #555; font-size: 14px; line-height: 1.6;">
                    💡 <strong>Tip:</strong> Love the content? Share Creator.ai with your colleagues and friends!
                </p>
                
                <p style="color: #555; font-size: 16px; line-height: 1.6; margin-bottom: 0;">
                    Best regards,<br>
                    🚀 <strong>The Creator.ai Team</strong><br>
                    <span style="color: #888; font-size: 14px;">Powered by LangGraph & Ollama</span>
                </p>
            </div>
            
            <div style="text-align: center; margin-top: 20px; color: #888; font-size: 12px;">
                <p>🤖 This email was generated and sent by Creator.ai</p>
            </div>
        </div>
        """,
        "attachments": [{
            "filename": filename,
            "content": docx_base64
        }]
    })


# -------------------------------
# Pipeline Runner
# -------------------------------
async def run_content_pipeline_async(input_data: dict) -> dict:
    """Run the full content creation pipeline"""
    state = ContentState(**input_data)
    
    # 7. ADD SEARCH AS FIRST STEP IN PIPELINE
    state = await web_search_async(state)
    state = await content_strategist_async(state)
    state = await content_writer_async(state)
    state = await content_reviewer_async(state)
    state = await finalizer_async(state)
    
    return state


# -------------------------------
# Authentication Functions
# -------------------------------
def verify_token(token: str) -> bool:
    return token in active_tokens

def authenticate_user(username: str, password: str) -> bool:
    return HARDCODED_USERS.get(username) == password

# -------------------------------
# API Endpoints
# -------------------------------
@app.post("/api/login", response_model=LoginResponse)
async def login(request: LoginRequest):
    """Login endpoint with hardcoded credentials"""
    if authenticate_user(request.username, request.password):
        token = secrets.token_urlsafe(32)
        active_tokens.add(token)
        return LoginResponse(token=token, message="Login successful")
    else:
        raise HTTPException(status_code=401, detail="Invalid credentials")

@app.post("/api/logout")
async def logout(token: str):
    """Logout endpoint"""
    active_tokens.discard(token)
    return {"message": "Logged out successfully"}

@app.get("/")
async def root():
    """Health check endpoint"""
    return {
        "message": "AI Content Creation Agent API",
        "status": "running",
        "model": LLM_MODEL,
        "endpoints": {
            "generate": "/api/generate",
            "stream": "/api/stream",
            "health": "/health"
        }
    }


@app.get("/health")
async def health_check():
    """Check if Ollama is running"""
    try:
        test_response = llm.invoke("Hi")
        return {"status": "healthy", "model": LLM_MODEL, "ollama": "connected"}
    except Exception as e:
        raise HTTPException(status_code=503, detail=f"Ollama not available: {str(e)}")


@app.post("/api/generate", response_model=ContentResponse)
async def generate_content(request: ContentRequest, authorization: str = Header(None)):
    """
    Generate content (non-streaming)
    Returns complete result when finished
    """
    # Check authentication
    if not authorization or not authorization.startswith("Bearer "):
        raise HTTPException(status_code=401, detail="Missing or invalid authorization header")
    
    token = authorization.split(" ")[1]
    if not verify_token(token):
        raise HTTPException(status_code=401, detail="Invalid or expired token")
    
    try:
        input_data = {
            "project_brief": request.project_brief,
            "audience": request.audience,
            "tone": request.tone,
            "keywords": request.keywords,
            "content_type": request.content_type,
            # 8. PASS WEB_SEARCH PARAM FROM REQUEST
            "web_search": request.web_search
        }
        
        result = await run_content_pipeline_async(input_data)
        
        return ContentResponse(
            status="completed" if not result.get("needs_escalation") else "escalated",
            strategy=result.get("strategy"),
            draft_content=result.get("draft_content"),
            quality_score=result.get("quality_score"),
            reviewer_comments=result.get("reviewer_comments"),
            needs_escalation=result.get("needs_escalation", False),
            final_content=result.get("final_content"),
            timestamp=datetime.now().isoformat()
        )
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@app.post("/api/stream")
async def stream_content(request: ContentRequest, authorization: str = Header(None)):
    """
    Stream content generation in real-time
    Returns Server-Sent Events (SSE)
    """
    # Check authentication
    if not authorization or not authorization.startswith("Bearer "):
        raise HTTPException(status_code=401, detail="Missing or invalid authorization header")
    
    token = authorization.split(" ")[1]
    if not verify_token(token):
        raise HTTPException(status_code=401, detail="Invalid or expired token")
    
    async def generate_stream():
        try:
            input_data = {
                "project_brief": request.project_brief,
                "audience": request.audience,
                "tone": request.tone,
                "keywords": request.keywords,
                "content_type": request.content_type,
                # 9. PASS WEB_SEARCH PARAM FROM REQUEST
                "web_search": request.web_search
            }
            
            state = ContentState(**input_data)
            
            # 10. ADD "STAGE 0: WEB SEARCH" AND FIX THE BUG
            yield f"data: {json.dumps({'stage': 'search', 'content': 'Searching the web...', 'progress': 10})}\n\n"
            state = await web_search_async(state)
            # THIS IS THE BUG FIX: Send 'search_results', not 'web_search'
            yield f"data: {json.dumps({'stage': 'search', 'content': state.get('search_results'), 'progress': 10})}\n\n"
            
            # Stage 1: Strategy
            yield f"data: {json.dumps({'stage': 'strategy', 'content': 'Generating strategy...', 'progress': 25})}\n\n"
            state = await content_strategist_async(state)
            yield f"data: {json.dumps({'stage': 'strategy', 'content': state.get('strategy'), 'progress': 25})}\n\n"
            
            # Stage 2: Writing
            yield f"data: {json.dumps({'stage': 'writing', 'content': 'Writing content...', 'progress': 50})}\n\n"
            state = await content_writer_async(state)
            yield f"data: {json.dumps({'stage': 'writing', 'content': state.get('draft_content'), 'progress': 50})}\n\n"
            
            # Stage 3: Review
            yield f"data: {json.dumps({'stage': 'review', 'content': 'Reviewing content...', 'progress': 75})}\n\n"
            state = await content_reviewer_async(state)
            review_data = {
                'score': state.get('quality_score'),
                'comments': state.get('reviewer_comments')
            }
            yield f"data: {json.dumps({'stage': 'review', 'content': json.dumps(review_data), 'progress': 75})}\n\n"
            
            # Stage 4: Finalize
            state = await finalizer_async(state)
            final_data = {
                'status': 'completed' if not state.get('needs_escalation') else 'escalated',
                'final_content': state.get('final_content'),
                'needs_escalation': state.get('needs_escalation')
            }
            yield f"data: {json.dumps({'stage': 'complete', 'content': json.dumps(final_data), 'progress': 100})}\n\n"
            
        except Exception as e:
            yield f"data: {json.dumps({'stage': 'error', 'content': str(e), 'progress': 0})}\n\n"
    
    return StreamingResponse(generate_stream(), media_type="text/event-stream")


@app.get("/api/models")
async def list_models():
    """List available Ollama models"""
    return {
        "current_model": LLM_MODEL,
        "recommended_models": [
            "llama3",
            "llama3:8b-instruct-q4_K_M",
            "phi3:mini",
            "gemma:2b",
            "mistral:7b-instruct"
        ]
    }

@app.post("/api/send-email")
async def send_email_endpoint(request: dict):
    """
    Send generated content via email
    """
    try:
        recipient_email = request.get("email")
        docx_base64 = request.get("docx_base64")  # Pre-generated DOCX from frontend
        project_brief = request.get("project_brief", "Generated Content")
        
        if not recipient_email or not docx_base64:
            raise HTTPException(status_code=400, detail="Email and DOCX content are required")
        
        await send_content_email_with_docx(recipient_email, docx_base64, project_brief)
        
        return {"status": "success", "message": "Email sent successfully!"}
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Failed to send email: {str(e)}")


@app.get("/api/download-image")
async def download_image(url: str, request: ContentRequest):
    """
    Proxies an image download to bypass CORS.
    """
    if not url:
        raise HTTPException(status_code=400, detail="No URL provided")

    try:
        async with httpx.AsyncClient() as client:
            # Stream the request to the external URL
            response = await client.stream("GET", url, headers={"User-Agent": request.headers.get("User-Agent", "FastAPI-Proxy")})
            
            # Check if the response is successful and an image
            if response.status_code != 200:
                raise HTTPException(status_code=response.status_code, detail="Failed to fetch image")
            
            content_type = response.headers.get("content-type", "application/octet-stream")
            if not content_type.startswith("image/"):
                raise HTTPException(status_code=400, detail="URL does not point to a valid image")

            # Get filename from URL
            filename = url.split('/')[-1].split('?')[0] or "image.jpg"
            
            return StreamingResponse(
                response.aiter_bytes(), 
                media_type=content_type,
                headers={"Content-Disposition": f"attachment; filename=\"{filename}\""}
            )

    except httpx.RequestError as e:
        raise HTTPException(status_code=500, detail=f"Error fetching image: {e}")


# -------------------------------
# Run Server
# -------------------------------
if __name__ == "__main__":
    import uvicorn
    print("🚀 Starting AI Content Creation Agent API...")
    print("📝 Docs available at: http://localhost:8000/docs")
    print("🔄 Stream endpoint: http://localhost:8000/api/stream")
    print("⚡ Generate endpoint: http://localhost:8000/api/generate")
    uvicorn.run(app, host="0.0.0.0", port=8000, reload=True)